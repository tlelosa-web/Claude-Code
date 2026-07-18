import docx
import openpyxl
import os
import sys

project_root = os.path.dirname(os.path.abspath(__file__))
base = os.path.join(project_root, '2_Source_Data')

docx_files = [
    'Curriculum Vitae Tebello 052014.docx',
    'Curriculum Vitae Tebello 2020.docx',
    'CV - TL 2021-11.docx',
    'CV Template - TL 2021.docx',
]

xlsx_files = [
    'CV_Tebello.xlsx',
    'Copy of CV_Tebello2017.xlsx',
    'CV_Tebello2020.xlsx',
]

out_path = os.path.join(project_root, 'extracted_content.txt')
out = open(out_path, 'w', encoding='utf-8')

def write(text):
    out.write(text + '\n')
    sys.stdout.write(text + '\n')

for fname in docx_files:
    fpath = os.path.join(base, fname)
    write('=' * 80)
    write(f'FILE: {fname}')
    write('=' * 80)
    try:
        doc = docx.Document(fpath)
        write(f'Paragraphs: {len(doc.paragraphs)}')
        write(f'Tables: {len(doc.tables)}')
        write('')
        write('--- PARAGRAPHS ---')
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                style = para.style.name if para.style else 'None'
                write(f'[{style}] {para.text}')
        write('')
        for i, table in enumerate(doc.tables):
            write(f'--- TABLE {i+1} ---')
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                write(' | '.join(cells))
            write('')
    except Exception as e:
        write(f'ERROR: {e}')
    write('')

for fname in xlsx_files:
    fpath = os.path.join(base, fname)
    write('=' * 80)
    write(f'FILE: {fname}')
    write('=' * 80)
    try:
        wb = openpyxl.load_workbook(fpath)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            write(f'--- SHEET: {sheet_name} (rows={ws.max_row}, cols={ws.max_column}) ---')
            for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column, values_only=False):
                vals = []
                for cell in row:
                    if cell.value is not None:
                        vals.append(f'{cell.coordinate}={cell.value}')
                if vals:
                    write(' | '.join(str(v) for v in vals))
            write('')
    except Exception as e:
        write(f'ERROR: {e}')
    write('')

write('DONE')
out.close()
