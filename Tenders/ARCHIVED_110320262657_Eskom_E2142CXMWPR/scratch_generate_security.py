import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_security_docs():
    basedir = r"c:\Users\tlelo\Desktop\110320262657\submission\technical"
    os.makedirs(basedir, exist_ok=True)
    
    # 1. Network Security Design
    out_path1 = os.path.join(basedir, "E2142CXMWPR_Network_Security_Design_v1.0_Delta_20260313.docx")
    doc1 = Document()
    title1 = doc1.add_heading('Network Security Design and Rules Matrix', 0)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc1.add_paragraph('Eskom GIT Enterprise Historian System\nPrepared by: Delta (Network & Security Engineer)\nDate: 2026-03-13\n')
    
    doc1.add_heading('1. Network Segmentation Strategy', level=1)
    doc1.add_paragraph('The Enterprise Historian network topography leverages an IT/OT DMZ model based on Purdue Enterprise Reference Architecture (PERA) to strictly enforce segmentation. Firewalls reside between OT, DMZ, and IT zones. Data collectors operate in OT at level 3, pushing data securely through the DMZ (Level 3.5) to the IT Historian (Level 4).')
    
    doc1.add_heading('2. Firewall Rules Matrix', level=1)
    doc1.add_paragraph('All firewall rules default to implicit deny. Required communication ports are limited to necessary protocols (OPC UA, MQTT, HTTPS).')
    table1 = doc1.add_table(rows=1, cols=5)
    table1.style = 'Table Grid'
    hdr1 = table1.rows[0].cells
    hdr1[0].text = 'Source Zone'
    hdr1[1].text = 'Dest Zone'
    hdr1[2].text = 'Protocol'
    hdr1[3].text = 'Port'
    hdr1[4].text = 'Purpose'
    
    rules = [
        ("OT (L3)", "DMZ (L3.5)", "TCP (OPC UA)", "4840", "Collector to DMZ Buffer node"),
        ("DMZ (L3.5)", "IT (L4)", "TCP (HTTPS)", "443", "Buffer node to Central Historian"),
        ("IT (L4)", "DMZ (L3.5)", "TCP/UDP (NTP)", "123", "Time Synchronization"),
        ("IT (L4)", "DMZ (L3.5)", "TCP (LDAPS)", "636", "AD Authentication")
    ]
    for src, dst, proto, port, purp in rules:
        row = table1.add_row().cells
        row[0].text = src
        row[1].text = dst
        row[2].text = proto
        row[3].text = port
        row[4].text = purp

    doc1.add_heading('3. VPN and Remote Access', level=1)
    doc1.add_paragraph('Remote vendor access is restricted via bastion hosts. All VPN sessions require Multi-Factor Authentication (MFA) and are strictly monitored with automatic session timeouts.')
    doc1.save(out_path1)
    print(f"File successfully created: {out_path1}")

    # 2. Security Hardening Checklist
    out_path2 = os.path.join(basedir, "E2142CXMWPR_Security_Hardening_Checklist_v1.0_Delta_20260313.docx")
    doc2 = Document()
    title2 = doc2.add_heading('Security Hardening Checklist & Evidence Pack', 0)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc2.add_paragraph('Eskom GIT Enterprise Historian System\nPrepared by: Delta (Network & Security Engineer)\nDate: 2026-03-13\n')
    
    doc2.add_heading('1. TLS Hardening and PKI Integration', level=1)
    doc2.add_paragraph('All communications are encrypted using TLS 1.2 or higher. Certificates are provisioned by an internal Enterprise Root CA.')
    
    doc2.add_heading('2. Server and Endpoint Hardening Baseline (CIS Benchmarks)', level=1)
    doc2.add_paragraph('Checklist:')
    hardening_items = [
        "Disable unneeded services (e.g., FTP, Telnet).",
        "Enforce complex password policies via GPO.",
        "Enable local OS firewalls with default-block rules.",
        "Uninstall all extraneous applications and utilities.",
        "Implement Endpoint Detection and Response (EDR) agents."
    ]
    for item in hardening_items:
        doc2.add_paragraph(item, style='List Bullet')
        
    doc2.add_heading('3. Application Security Configuration', level=1)
    doc2.add_paragraph('Historian Service accounts utilize Managed Service Accounts (gMSA) or apply the principle of least privilege, with passwords rotated every 60 days where applicable.')
    doc2.save(out_path2)
    print(f"File successfully created: {out_path2}")

    # 3. Audit/Monitoring Plan
    out_path3 = os.path.join(basedir, "E2142CXMWPR_Audit_Monitoring_Plan_v1.0_Delta_20260313.docx")
    doc3 = Document()
    title3 = doc3.add_heading('Audit & Monitoring Plan', 0)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc3.add_paragraph('Eskom GIT Enterprise Historian System\nPrepared by: Delta (Network & Security Engineer)\nDate: 2026-03-13\n')
    
    doc3.add_heading('1. Centralized Logging Architecture', level=1)
    doc3.add_paragraph('Security logs from Historian servers, DMZ proxies, and firewalls are forwarded via syslog to a central SIEM platform. Log retention is set to 1 year hotline/cold-line as per Eskom policy.')
    
    doc3.add_heading('2. Monitored Events and Alerting Configuration', level=1)
    doc3.add_paragraph('Events configured for immediate alert generation:')
    alerts = [
        "Multiple failed login attempts (Brute force detection).",
        "Privilege escalation anomalies.",
        "Firewall rule block threshold exceeded.",
        "Unexpected collector connection drop.",
        "Certificate expiration within 30 days."
    ]
    for a in alerts:
        doc3.add_paragraph(a, style='List Bullet')
        
    doc3.add_heading('3. Audit Cadence', level=1)
    doc3.add_paragraph('Security assessments and vulnerability scanning will be performed monthly. Identified CVEs with CVSS scores > 7.0 will be prioritized and mediated strictly within 14 days.')
    
    doc3.save(out_path3)
    print(f"File successfully created: {out_path3}")

if __name__ == '__main__':
    create_security_docs()
