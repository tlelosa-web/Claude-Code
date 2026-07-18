import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contractual_response():
    # Setup directory
    basedir = r"c:\Users\tlelo\Desktop\110320262657\submission\contracts"
    os.makedirs(basedir, exist_ok=True)
    out_path = os.path.join(basedir, "E2142CXMWPR_Contractual_Response_v1.0_Juliet_20260313.docx")

    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Contractual Response & Legal Review', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Eskom GIT Enterprise Historian Licence, Maintenance & Support (ITT E2142CXMWPR)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = subtitle.runs[0]
    runner.font.size = Pt(14)
    runner.font.bold = True
    
    # Author
    doc.add_paragraph('\nPrepared by: Juliet (Contracts and Legal Lead)\nDate: 2026-03-13\n')

    # Section 1: NEC3 Review Memo
    doc.add_heading('1. NEC3 Professional Services Contract (PSC) Review Memo', level=1)
    doc.add_paragraph(
        "This memorandum summarizes the contractual framework under the proposed NEC3 PSC for the Enterprise "
        "Historian project. As an off-the-shelf software licensing, maintenance, and support engagement, the "
        "rigid structure of the NEC3 PSC must be carefully managed, particularly regarding Compensation Events (CEs) "
        "and Early Warnings (EWs)."
    )
    doc.add_paragraph(
        "Key aspects noted from the review of Annexure L and Annexure T:\n"
        "• Early Warnings: Both parties must notify as soon as a matter arises that could increase the total of the Prices, "
        "delay Completion, or impair the performance of the service in use.\n"
        "• Compensation Events: Limited to those defined in the core clauses. We have noted the necessity of strict "
        "adherence to timelines for notifying CEs (usually within 8 weeks) to preserve entitlement to time and cost adjustments.\n"
        "• Accepted Programme: The implementation schedule and SLAs will form part of the Accepted Programme, requiring "
        "regular updating and formal acceptance by the Employer."
    )

    # Section 2: Contractual Risk Register
    doc.add_heading('2. Contractual Risk Register', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Risk Description'
    hdr_cells[1].text = 'Contractual Impact'
    hdr_cells[2].text = 'Proposed Mitigation / Clause Reference'
    
    risks = [
        ("Indemnity for Indirect / \nConsequential Loss", "High exposure if not capped. Software deployment carries risk of data/production loss.", "Insert a Z-clause limiting liability for consequential loss and capping overall liability to the contract value."),
        ("Intellectual Property (IP) Ownership", "Standard NEC3 may favor Employer ownership. Product is proprietary COTS software.", "Standard Z-clause clarification ensuring pre-existing IP and COTS software IP remains with the Supplier."),
        ("Software Defects vs. 'Defects' under PSC", "Correction of Defects provisions in standard PSC are misaligned with SaaS SLA remediation timelines.", "Qualify the Definition of 'Defects' to refer to Software SLA severity matrix instead of standard project works defects.")
    ]
    for risk, impact, mitigation in risks:
        row_cells = table.add_row().cells
        row_cells[0].text = risk
        row_cells[1].text = impact
        row_cells[2].text = mitigation

    # Section 3: Deviation & Qualification List
    doc.add_heading('3. Deviation and Qualification List (Z-Clauses Proposed)', level=1)
    doc.add_paragraph(
        "The Supplier accepts the standard NEC3 PSC and Eskom Conditions of Tender subject to the following proposed qualifications:"
    )
    doc.add_paragraph("1. Limitation of Liability (Z-Clause Amendment): ", style='List Number').add_run("The Supplier's total aggregate liability under this agreement, whether in contract, delict or otherwise, shall not exceed 100% of the total contract value. Neither party shall be liable for any indirect, special, or consequential loss, including loss of profits, data, or production.")
    doc.add_paragraph("2. Intellectual Property Rights: ", style='List Number').add_run("Notwithstanding any core clause to the contrary, the Supplier retains all ownership and Intellectual Property Rights in its Pre-existing Background IP and standard commercially available software (Enterprise Historian). The Employer is granted a license as per the EULA.")
    doc.add_paragraph("3. Performance Measurement: ", style='List Number').add_run("Performance deductions and Defect correction timelines shall be governed solely by the agreed Service Level Agreement (SLA) rather than standard NEC3 Defect Correction periods.")

    # Section 4: Contractual Compliance Summary
    doc.add_heading('4. Contractual Compliance Summary', level=1)
    doc.add_paragraph(
        "Except as expressly set out in the Deviation and Qualification List above, the Supplier confirms full "
        "compliance with the Standard Conditions of Tender (Annexure T) and the proposed NEC3 Professional "
        "Services Contract (Annexure L). The Supplier confirms its readiness to participate in formal contract "
        "negotiations based on these stated qualifications, and asserts that these standard software industry "
        "qualifications do not materially alter the risk profile to Eskom."
    )

    doc.add_page_break()
    p = doc.add_paragraph('--- END OF MEMO ---')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"File successfully created at: {out_path}")

if __name__ == '__main__':
    create_contractual_response()
