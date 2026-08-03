import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_service_plan():
    # Setup directory
    basedir = r"c:\Users\tlelo\Desktop\110320262657\submission\service"
    os.makedirs(basedir, exist_ok=True)
    out_path = os.path.join(basedir, "E2142CXMWPR_Support_SLA_Plan_v1.0_Foxtrot_20260313.docx")

    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('Service Support & Maintenance SLA Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Eskom GIT Enterprise Historian Licence, Maintenance & Support (ITT E2142CXMWPR)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = subtitle.runs[0]
    runner.font.size = Pt(14)
    runner.font.bold = True
    
    # Author
    doc.add_paragraph('\nPrepared by: Foxtrot (Service Delivery Lead)\nDate: 2026-03-13\n')

    # Section 1: Support Model
    doc.add_heading('1. Service Support Model', level=1)
    doc.add_paragraph('This document defines the ITIL-aligned support framework and Service Level Agreement (SLA) for the Enterprise Historian system. It ensures high availability and efficient incident resolution.')
    doc.add_paragraph('Service Hours:', style='List Bullet')
    doc.add_paragraph('Business Hours (L1 & L2): 08:00 - 17:00 (SAST) Monday - Friday.', style='List Continue')
    doc.add_paragraph('After-Hours / 24x7 (L3 & Emergency Incident): Escrow response on critical P1 issues.', style='List Continue')
    doc.add_paragraph('Support Tiers:', style='List Bullet')
    doc.add_paragraph('Tier 1 (Service Desk): First point of contact, triage, and basic troubleshooting.', style='List Continue')
    doc.add_paragraph('Tier 2 (Application Support): Advanced diagnosis, configuration issues, data historian queries.', style='List Continue')
    doc.add_paragraph('Tier 3 (OEM/Engineering): Bug fixes, patching, critical system down restorations.', style='List Continue')

    # Section 2: KPIs & SLAs
    doc.add_heading('2. Key Performance Indicators (KPIs) & SLAs', level=1)
    doc.add_paragraph('The following SLA targets apply to all registered incidents and problems:')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity Level'
    hdr_cells[1].text = 'Response Time'
    hdr_cells[2].text = 'Target Resolution'
    hdr_cells[3].text = 'System Availability Target'
    
    slas = [
        ("P1 - Critical (System Down)", "15 Minutes", "4 Hours", "99.9%"),
        ("P2 - High (Major Degradation)", "30 Minutes", "8 Hours", "99.9%"),
        ("P3 - Medium (Minor Degradation)", "4 Hours", "2 Business Days", "N/A"),
        ("P4 - Low (Informational/Request)", "8 Hours", "5 Business Days", "N/A")
    ]
    for sev, resp, res, avail in slas:
        row_cells = table.add_row().cells
        row_cells[0].text = sev
        row_cells[1].text = resp
        row_cells[2].text = res
        row_cells[3].text = avail
        
    doc.add_paragraph('Operational Level Agreements (OLAs) with internal networking, security, and infrastructure teams will align to support these targets. Maintenance windows are scheduled for Sunday 12:00 AM - 04:00 AM (SAST).')

    # Section 3: ITIL Processes (Incident, Problem, Change)
    doc.add_heading('3. ITIL Processes Integration', level=1)
    doc.add_paragraph('Incident Management:', style='List Bullet')
    doc.add_paragraph('All incidents are logged in the ITSM tool (ServiceNow or equivalent) and mapped to the configuration item (CI) for the Enterprise Historian. Escalations follow the defined severity and response matrix.', style='List Continue')

    doc.add_paragraph('Problem Management:', style='List Bullet')
    doc.add_paragraph('For recurring P2s or any P1 incident, a formal Root Cause Analysis (RCA) is required within 5 business days, followed by entry into the Known Error Database (KEDB).', style='List Continue')
    
    doc.add_paragraph('Change Management:', style='List Bullet')
    doc.add_paragraph('All deployments, patches, and version upgrades require Cab (Change Advisory Board) approval. Emergency Changes for P1 restoration bypass CAB but require Post-Implementation Review (PIR).', style='List Continue')

    # Section 4: Escalation & Governance
    doc.add_heading('4. Escalation Matrix and Governance', level=1)
    doc.add_paragraph('A hierarchical escalation path is established for SLA breaches or impending risks:')
    doc.add_paragraph('L1/L2 Engineer -> Service Delivery Manager -> Project Executive -> Eskom IT/OT Director')
    
    doc.add_paragraph('Reporting Frequency:', style='List Bullet')
    doc.add_paragraph('Weekly: Incident volume, SLA breaches, patching compliance.', style='List Continue')
    doc.add_paragraph('Monthly: Service Review Meeting (Major incidents, capacity, availability).', style='List Continue')
    doc.add_paragraph('Quarterly: Continual Service Improvement (CSI) strategy assessment.', style='List Continue')

    doc.add_page_break()
    p = doc.add_paragraph('--- END OF SLA PLAN ---')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"File successfully created at: {out_path}")

if __name__ == '__main__':
    create_service_plan()
