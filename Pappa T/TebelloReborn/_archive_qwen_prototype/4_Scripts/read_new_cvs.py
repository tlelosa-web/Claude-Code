from docx import Document
from pathlib import Path
import sys

project_root = Path(__file__).resolve().parents[1]

files = [
    project_root / '2_Source_Data' / 'Tebello Lelosa CV - 2025 new.docx',
    project_root / '2_Source_Data' / 'Tebello Lelosa CV .docx',
]

for f in files:
    print('=' * 80)
    print(f'FILE: {f.name}')
    print('=' * 80)
    try:
        doc = Document(f)
        print(f'Paragraphs: {len(doc.paragraphs)}')
        print(f'Tables: {len(doc.tables)}')
        for p in doc.paragraphs:
            if p.text.strip():
                print(f'[{p.style.name}] {p.text}')
        for ti, t in enumerate(doc.tables):
            print(f'\n--- TABLE {ti+1} ---')
            for row in t.rows:
                print(' | '.join(c.text.strip() for c in row.cells))
    except Exception as e:
        print(f'Error: {e}', file=sys.stderr)
    print()
